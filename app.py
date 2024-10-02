from flask import Flask, render_template
from markupsafe import escape
from docx import Document

app = Flask(__name__)

def print_docx_content(docx_path):
    doc = Document(docx_path)
    
    fulltext = []

    for paragraph in doc.paragraphs:
        fulltext.append(paragraph.text)
    return "erm"
    return '\n'.join(fulltext)

@app.route('/')
@app.route('/index/')
def index():

    document = Document("Sorrows of Empire.docx")
    fulltext = []

    for paragraph in document.paragraphs:
        fulltext.append(paragraph.text)
    itemsToWriteToFile = itemsToWriteToFile.translate(None, "(),\"\\n")
    return ''.join(fulltext)

    return '\n'.join(fulltext)
#    print_docx_content("Sorrows of Empire.docx")
    return "test"


@app.route('/about/')
def about():
    return '<h3>This is a Flask web application.</h3>'